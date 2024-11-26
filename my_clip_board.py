from argparse import ArgumentParser, Namespace
from io import BytesIO
import json
import keyboard as k
import os
from PIL import Image, ImageGrab
import shelve
import sys
import time

from docx import Document
from docx.shared import Pt
import pyttsx3
import pyperclip
import win32clipboard

from mcb.vars.paths import FILEPATH_OF_MCB_LIST_WORD_DOC, FILEPATH_OF_MCB_RECORD_WORD_DOC, FILEPATH_OF_MCB_IMG_WORD_DOC, TO_FOLDER, TO_RECORD

class MyClipBoard:
    def __init__(self) -> None:
        self.myShelf = shelve.open(fr'{TO_FOLDER}')
        self.recordShelf = shelve.open(fr'{TO_RECORD}')
        if "records" not in self.recordShelf:
            self.recordShelf["records"] = []

    def play_voice(
        self,
        args: Namespace,
        read_this: str
    ):
        if args.unvoiced == True:
            return
        
        self.engine = pyttsx3.init()
        self.voices = self.engine.getProperty(name="voices")
        self.engine.setProperty(
            name="voice",
            value=self.voices[3].id
        )
        self.engine.setProperty(
            name="rate",
            value=180
        )
        self.engine.say(
            text=read_this
        )
        self.engine.runAndWait()

    def parse_args(
        self,
        manual_args: list = None
    ):
        self.parser = ArgumentParser(
            exit_on_error=False
        )
        self.parser.add_argument(
            "keyword",
            nargs="?",
            default=None
        )
        self.parser.add_argument(
            "-s",
            "--save",
            action="store_true",
            default=False
        )
        self.parser.add_argument(
            "-d",
            "--delete",
            action="store_true",
            default=False
        )
        self.parser.add_argument(
            "-uv",
            "--unvoiced",
            action="store_true",
            default=False
        )
        self.parser.add_argument(
            "-sh",
            "--show",
            action="store_true",
            default=False
        )
        self.parser.add_argument(
            "-jkv",
            "--json_key_value",
            action="store_true",
            default=False
        )
        self.parser.add_argument(
            "-jst",
            "--json_structure",
            action="store_true",
            default=False
        )
        self.parser.add_argument(
            "-tl",
            "--tab_length",
            type=int,
            default=None
        )
        self.parser.add_argument(
            "-tc",
            "--tab_count",
            type=int,
            default=None
        )
        self.parser.add_argument(
            "-l",
            "--list",
            action="store_true",
            default=False
        )
        self.parser.add_argument(
            "-r",
            "--record",
            action="store_true",
            default=False
        )
        self.parser.add_argument(
            "-dr",
            "--delete_record",
            action="store_true",
            default=False
        )
        
        parsed_args = self.parser.parse_args() if manual_args == None else self.parser.parse_args(args=manual_args)
        
        return parsed_args
    
    def save(
        self,
        args: Namespace
    ):
        if isinstance(ImageGrab.grabclipboard(), Image.Image):
            with self.myShelf as myShelf:
                myShelf[args.keyword] = ImageGrab.grabclipboard()
                self.play_voice(
                    args=args,
                    read_this="イメージを保存した"
                )
        elif isinstance(ImageGrab.grabclipboard(), list):
            strPath = ImageGrab.grabclipboard()[0]
            
            with self.myShelf as myShelf:
                myShelf[args.keyword] = Image.open(strPath)
                self.play_voice(
                    args=args,
                    read_this="イメージを保存した"
                )
        else:
            with self.myShelf as myShelf:
                myShelf[args.keyword] = pyperclip.paste()
                self.play_voice(
                    args=args,
                    read_this="なんかを保存した"
                )

    def load(
        self,
        args: Namespace
    ):
        with self.myShelf as myShelf:
            if isinstance(myShelf[args.keyword], Image.Image):
                self.send_image_to_clipboard(
                    image=myShelf[args.keyword]
                )
                self.play_voice(
                    args=args,
                    read_this="イメージをクリップボードにコピーした"
                )
            
            else:
                copyMe = myShelf[args.keyword]
                if args.json_key_value or args.json_structure:
                    copyMe = self.jsonPrep(
                        args=args,
                        prepThis=copyMe
                    )
                pyperclip.copy(copyMe)
                self.play_voice(
                    args=args,
                    read_this="なんかをクリップボードにコピーした"
                )

    def list_keywords(self): 
        with self.myShelf as myShelf:
            toCopy = '\n'.join(list(myShelf.keys()))
            doc = Document()
            para = doc.add_paragraph()
            run = para.add_run(
                text=toCopy
            )
            font = run.font
            font.name = 'Calibri'
            font.size = Pt(15)
            doc.save(str(FILEPATH_OF_MCB_LIST_WORD_DOC))
            
            os.system(r"start WINWORD.exe " + str(FILEPATH_OF_MCB_LIST_WORD_DOC))

    def list_records(self): 
        with self.recordShelf as recordShelf:
            argsRecord: list = recordShelf[list(recordShelf.keys())[0]]
            toCopy = '\n'.join(argsRecord)
            doc = Document()
            para = doc.add_paragraph()
            run = para.add_run(
                text=toCopy
            )
            font = run.font
            font.name = 'Calibri'
            font.size = Pt(15)
            doc.save(str(FILEPATH_OF_MCB_RECORD_WORD_DOC))
            
            os.system(r"start WINWORD.exe " + str(FILEPATH_OF_MCB_RECORD_WORD_DOC))

    def delete(
        self,
        args: Namespace
    ): 
        with self.myShelf as myShelf:
            del myShelf[args.keyword]

    def show_image(
        self,
        args: Namespace
    ):
        with self.myShelf as myShelf:
            if isinstance(myShelf[args.keyword], Image.Image):
                self.send_image_to_clipboard(myShelf[args.keyword])
                os.system(r"start WINWORD.exe " + str(FILEPATH_OF_MCB_IMG_WORD_DOC))
                time.sleep(2.5)
                k.press_and_release("ctrl+v")
                k.press_and_release("ctrl+s")
            else:
                self.play_voice(
                    args=args,
                    read_this="これがイメージじゃない"
                )
    
    def send_image_to_clipboard(
        self,
        image: Image.Image
    ):
        output = BytesIO()
        image.convert('RGB').save(output, 'BMP')
        data = output.getvalue()[14:]
        output.close()

        win32clipboard.OpenClipboard()
        win32clipboard.EmptyClipboard()
        win32clipboard.SetClipboardData(
            win32clipboard.CF_DIB,
            data
        )
        win32clipboard.CloseClipboard()

    def jsonPrep(
        self,
        args: Namespace,
        prepThis: str
    ):
        if (prepThis[0] == "{" and prepThis[-1] == "}") == False:
            prepThis = "{" + prepThis + "}"
        
        hold = json.loads(s=prepThis)
        hold = json.dumps(
            obj=hold,
            ensure_ascii=False,
            indent=args.tab_length
        )
        if args.json_key_value:
            hold = hold.strip("{").strip("}").strip()
        split_newline = hold.split("\n")
        whitespace = " " * int(args.tab_length) * int(args.tab_count - 1)
        whitespace = "\n" + whitespace
        hold = whitespace.join(split_newline)
        return hold
    
    def record_args(self):
        with self.recordShelf as recordShelf:
            argsRecord: list = recordShelf[list(recordShelf.keys())[0]]
            args_string = " ".join(sys.argv[1::])
            if args_string in argsRecord:
                argsRecord.remove(args_string)
            argsRecord.append(args_string)
            recordShelf[list(recordShelf.keys())[0]] = argsRecord

    def delete_record(
        self,
        args: Namespace
    ):
        with self.recordShelf as recordShelf:
            argsRecord: list = recordShelf[list(recordShelf.keys())[0]]

            blacklist = [
                "-dr",
                "--delete_record"
            ]

            dr_removed_args = [elem for elem in sys.argv[1::] if elem not in blacklist]
            dr_removed_args_str = " ".join(dr_removed_args)
            
            if dr_removed_args_str in argsRecord:
                argsRecord.remove(dr_removed_args_str)
                print("deleted a record")
            else:
                raise ValueError(f"given input '{dr_removed_args_str}' not found in records")
            
            recordShelf[list(recordShelf.keys())[0]] = argsRecord

    def decision_tree(
        self,
        args: Namespace
    ):
        if args.keyword == None:
            if args.list:
                self.list_keywords()
                return
            
            elif args.record:
                self.list_records()
                return
            
            elif args.delete_record:
                self.delete_record(args=args)
                return
        
        else:
            if args.save:
                self.save(args=args)

            elif args.delete:
                self.delete(args=args)

            elif args.show:
                self.show_image(args=args)

            elif args.delete_record:
                self.delete_record(args=args)
                return
        
            # only keyword in args
            else: 
                self.load(args=args)

            self.record_args()